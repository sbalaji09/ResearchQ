"""
Literature Review Generator

Generates structured academic literature reviews from a collection of research papers.
Supports multiple citation styles (APA, MLA, Chicago) and export formats (Markdown, LaTeX, Word).
"""

import os
from dataclasses import dataclass, field
from datetime import datetime, timezone
from typing import Dict, List, Optional, Any
from pathlib import Path

from openai import OpenAI

from literature_review import (
    get_abstract,
    get_methodology,
    get_conclusion,
    get_paper_chunks,
    compare_papers,
    extract_methodology_summary,
)
from paper_store import paper_store

# Lazy initialization of OpenAI client
_openai_client = None

def get_openai_client():
    global _openai_client
    if _openai_client is None:
        _openai_client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY"))
    return _openai_client


# ============================================================================
# Data Classes
# ============================================================================

@dataclass
class PaperReference:
    """Represents a paper with citation metadata."""
    pdf_id: str
    title: str
    authors: List[str] = field(default_factory=list)
    year: Optional[int] = None
    abstract: Optional[str] = None

    def to_dict(self) -> Dict[str, Any]:
        return {
            "pdf_id": self.pdf_id,
            "title": self.title,
            "authors": self.authors,
            "year": self.year,
            "abstract": self.abstract,
        }


@dataclass
class LiteratureReviewResult:
    """The generated literature review with all sections."""
    title: str
    introduction: str
    methodology_overview: str
    key_findings: str
    research_gaps: str
    conclusion: str
    references: List[Dict[str, Any]]

    # Metadata
    papers_analyzed: List[str] = field(default_factory=list)
    citation_style: str = "apa"
    created_at: str = field(default_factory=lambda: datetime.now(timezone.utc).isoformat())

    def to_dict(self) -> Dict[str, Any]:
        return {
            "title": self.title,
            "introduction": self.introduction,
            "methodology_overview": self.methodology_overview,
            "key_findings": self.key_findings,
            "research_gaps": self.research_gaps,
            "conclusion": self.conclusion,
            "references": self.references,
            "papers_analyzed": self.papers_analyzed,
            "citation_style": self.citation_style,
            "created_at": self.created_at,
        }


# ============================================================================
# Citation Formatting
# ============================================================================

def format_citation_apa(paper: PaperReference, index: int) -> str:
    """Format a citation in APA style."""
    authors_str = ""
    if paper.authors:
        if len(paper.authors) == 1:
            authors_str = paper.authors[0]
        elif len(paper.authors) == 2:
            authors_str = f"{paper.authors[0]} & {paper.authors[1]}"
        else:
            authors_str = f"{paper.authors[0]} et al."
    else:
        authors_str = paper.title[:20]

    year = paper.year or "n.d."
    return f"({authors_str}, {year})"


def format_citation_mla(paper: PaperReference, index: int) -> str:
    """Format a citation in MLA style."""
    if paper.authors:
        author = paper.authors[0].split()[-1]  # Last name
    else:
        author = paper.title[:20]
    return f"({author})"


def format_citation_chicago(paper: PaperReference, index: int) -> str:
    """Format a citation in Chicago style (footnote number)."""
    return f"[{index}]"


def format_reference_apa(paper: PaperReference) -> str:
    """Format a full reference in APA style."""
    authors_str = ", ".join(paper.authors) if paper.authors else paper.title
    year = paper.year or "n.d."
    return f"{authors_str} ({year}). {paper.title}."


def format_reference_mla(paper: PaperReference) -> str:
    """Format a full reference in MLA style."""
    authors_str = ", ".join(paper.authors) if paper.authors else ""
    return f"{authors_str}. \"{paper.title}.\""


def format_reference_chicago(paper: PaperReference, index: int) -> str:
    """Format a full reference in Chicago style."""
    authors_str = ", ".join(paper.authors) if paper.authors else ""
    year = paper.year or "n.d."
    return f"[{index}] {authors_str}, \"{paper.title},\" {year}."


def get_citation_formatter(style: str):
    """Get the citation formatter for the given style."""
    formatters = {
        "apa": format_citation_apa,
        "mla": format_citation_mla,
        "chicago": format_citation_chicago,
    }
    return formatters.get(style.lower(), format_citation_apa)


def get_reference_formatter(style: str):
    """Get the reference formatter for the given style."""
    formatters = {
        "apa": format_reference_apa,
        "mla": format_reference_mla,
        "chicago": format_reference_chicago,
    }
    return formatters.get(style.lower(), format_reference_apa)


# ============================================================================
# Paper Data Collection
# ============================================================================

def collect_paper_data(pdf_ids: List[str]) -> List[PaperReference]:
    """Collect metadata and content for all papers."""
    papers = []

    for pdf_id in pdf_ids:
        # Get metadata from paper store
        stored_paper = paper_store.get_paper(pdf_id)

        title = pdf_id.replace("_", " ").replace("-", " ")
        authors = []
        year = None
        abstract = None

        if stored_paper:
            title = stored_paper.title or title
            authors = stored_paper.authors or []
            abstract = stored_paper.abstract

        # Try to get abstract from chunks if not in store
        if not abstract:
            abstract = get_abstract(pdf_id)
            if abstract:
                abstract = abstract[:500]

        papers.append(PaperReference(
            pdf_id=pdf_id,
            title=title,
            authors=authors,
            year=year,
            abstract=abstract,
        ))

    return papers


# ============================================================================
# Section Generators
# ============================================================================

def generate_introduction(papers: List[PaperReference], topic: Optional[str]) -> str:
    """Generate the introduction section."""
    abstracts_text = "\n\n".join([
        f"Paper: {p.title}\nAbstract: {p.abstract or 'Not available'}"
        for p in papers
    ])

    topic_context = f"on the topic of '{topic}'" if topic else ""

    prompt = f"""You are writing the introduction section of an academic literature review {topic_context}.

Based on these {len(papers)} research papers:

{abstracts_text}

Write an introduction that:
1. Introduces the research area and its importance
2. Provides context for why this topic matters
3. Briefly mentions the scope of papers being reviewed
4. States the purpose of the literature review

Write 2-3 paragraphs in formal academic style. Do not use citations yet - this is just the introduction."""

    response = get_openai_client().chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=800,
        temperature=0.4,
    )

    return response.choices[0].message.content.strip()


def generate_methodology_overview(pdf_ids: List[str], papers: List[PaperReference]) -> str:
    """Generate the methodology overview section."""
    # Get methodology summaries for each paper
    methodology_texts = []
    for pdf_id in pdf_ids:
        summary = extract_methodology_summary(pdf_id)
        if summary.get("summary") and "not found" not in summary["summary"].lower():
            paper = next((p for p in papers if p.pdf_id == pdf_id), None)
            title = paper.title if paper else pdf_id
            methodology_texts.append(f"**{title}**: {summary['summary']}")

    if not methodology_texts:
        return "Methodology information was not available for the reviewed papers."

    methodologies = "\n\n".join(methodology_texts)

    prompt = f"""You are writing the methodology overview section of an academic literature review.

Here are the methodologies used in the reviewed papers:

{methodologies}

Write a methodology overview that:
1. Summarizes the types of research designs used (qualitative, quantitative, mixed)
2. Compares sample sizes and participant characteristics
3. Discusses data collection methods
4. Notes any methodological patterns or trends

Write 2-3 paragraphs comparing and contrasting the methodological approaches.
Use phrases like "Several studies employed...", "In contrast, other researchers used...", etc."""

    response = get_openai_client().chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=800,
        temperature=0.4,
    )

    return response.choices[0].message.content.strip()


def generate_key_findings(pdf_ids: List[str], papers: List[PaperReference]) -> str:
    """Generate the key findings section with thematic organization."""
    # Collect content from papers
    all_content = []
    for i, pdf_id in enumerate(pdf_ids):
        abstract = get_abstract(pdf_id)
        conclusion = get_conclusion(pdf_id)
        paper = papers[i] if i < len(papers) else None
        title = paper.title if paper else pdf_id

        content = f"""
Paper [{i+1}]: {title}
Abstract: {abstract[:800] if abstract else 'Not available'}
Conclusions: {conclusion[:800] if conclusion else 'Not available'}
"""
        all_content.append(content)

    papers_text = "\n---\n".join(all_content)

    prompt = f"""You are writing the key findings section of an academic literature review.

Here are the papers with their abstracts and conclusions:

{papers_text}

Write a comprehensive findings section that:
1. Organizes findings THEMATICALLY, not paper-by-paper
2. Identifies 3-5 major themes across the papers
3. Uses citations like [1], [2], etc. to reference specific papers
4. Highlights where papers agree or reach similar conclusions
5. Notes any contradictions or debates between papers

Structure with subheadings for each theme. Write 4-5 paragraphs total.
Use formal academic language."""

    response = get_openai_client().chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=1500,
        temperature=0.4,
    )

    return response.choices[0].message.content.strip()


def generate_research_gaps(pdf_ids: List[str], papers: List[PaperReference]) -> str:
    """Generate the research gaps and future directions section."""
    # Collect limitations and future work mentions
    all_content = []
    for i, pdf_id in enumerate(pdf_ids):
        chunks = get_paper_chunks(pdf_id, section_filter=["limitation", "future", "conclusion", "discussion"])
        content = "\n".join([c["text"][:500] for c in chunks[:3]])
        paper = papers[i] if i < len(papers) else None
        title = paper.title if paper else pdf_id
        all_content.append(f"Paper: {title}\n{content}")

    papers_text = "\n---\n".join(all_content)

    prompt = f"""You are writing the research gaps section of an academic literature review.

Based on the limitations and discussion sections of these papers:

{papers_text}

Identify and write about:
1. Common limitations across the studies
2. Gaps in the current research (what hasn't been studied)
3. Methodological limitations
4. Populations or contexts that need more research
5. Suggested directions for future research

Write 2-3 paragraphs. Be specific about what is missing and why it matters.
This section should help researchers identify opportunities for new studies."""

    response = get_openai_client().chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=800,
        temperature=0.4,
    )

    return response.choices[0].message.content.strip()


def generate_conclusion(papers: List[PaperReference], key_findings: str) -> str:
    """Generate the conclusion section."""
    prompt = f"""You are writing the conclusion of an academic literature review that analyzed {len(papers)} research papers.

The key findings section covered:
{key_findings[:2000]}

Write a conclusion that:
1. Summarizes the main insights from the reviewed literature
2. Highlights the most important contributions
3. Provides a synthesis statement about the state of knowledge in this area
4. Ends with implications for practice or future research

Write 2 paragraphs in formal academic style."""

    response = get_openai_client().chat.completions.create(
        model="gpt-4o",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=600,
        temperature=0.4,
    )

    return response.choices[0].message.content.strip()


# ============================================================================
# Main Generator Function
# ============================================================================

def generate_literature_review(
    pdf_ids: List[str],
    topic: Optional[str] = None,
    citation_style: str = "apa",
) -> LiteratureReviewResult:
    """
    Generate a complete literature review from a collection of papers.

    Args:
        pdf_ids: List of PDF IDs to include in the review
        topic: Optional topic/title for the review
        citation_style: Citation style to use (apa, mla, chicago)

    Returns:
        LiteratureReviewResult with all sections
    """
    if len(pdf_ids) < 2:
        raise ValueError("Need at least 2 papers to generate a literature review")

    if len(pdf_ids) > 20:
        raise ValueError("Maximum 20 papers can be reviewed at once")

    # Collect paper data
    papers = collect_paper_data(pdf_ids)

    # Generate title
    if topic:
        title = f"Literature Review: {topic}"
    else:
        title = "Literature Review: A Synthesis of Current Research"

    # Generate each section
    introduction = generate_introduction(papers, topic)
    methodology_overview = generate_methodology_overview(pdf_ids, papers)
    key_findings = generate_key_findings(pdf_ids, papers)
    research_gaps = generate_research_gaps(pdf_ids, papers)
    conclusion = generate_conclusion(papers, key_findings)

    # Format references
    ref_formatter = get_reference_formatter(citation_style)
    references = []
    for i, paper in enumerate(papers):
        if citation_style == "chicago":
            ref_text = ref_formatter(paper, i + 1)
        else:
            ref_text = ref_formatter(paper)
        references.append({
            "index": i + 1,
            "pdf_id": paper.pdf_id,
            "formatted": ref_text,
        })

    return LiteratureReviewResult(
        title=title,
        introduction=introduction,
        methodology_overview=methodology_overview,
        key_findings=key_findings,
        research_gaps=research_gaps,
        conclusion=conclusion,
        references=references,
        papers_analyzed=pdf_ids,
        citation_style=citation_style,
    )


# ============================================================================
# Export Functions
# ============================================================================

def export_to_markdown(review: LiteratureReviewResult) -> str:
    """Export the literature review to Markdown format."""
    references_text = "\n".join([
        f"{r['index']}. {r['formatted']}"
        for r in review.references
    ])

    markdown = f"""# {review.title}

## Introduction

{review.introduction}

## Methodology Overview

{review.methodology_overview}

## Key Findings

{review.key_findings}

## Research Gaps and Future Directions

{review.research_gaps}

## Conclusion

{review.conclusion}

## References

{references_text}

---
*Generated on {review.created_at[:10]} | Citation style: {review.citation_style.upper()}*
"""
    return markdown


def export_to_latex(review: LiteratureReviewResult) -> str:
    """Export the literature review to LaTeX format."""
    # Escape special LaTeX characters
    def escape_latex(text: str) -> str:
        replacements = {
            '&': r'\&',
            '%': r'\%',
            '$': r'\$',
            '#': r'\#',
            '_': r'\_',
            '{': r'\{',
            '}': r'\}',
        }
        for char, replacement in replacements.items():
            text = text.replace(char, replacement)
        return text

    references_text = "\n".join([
        f"\\bibitem{{{r['pdf_id']}}} {escape_latex(r['formatted'])}"
        for r in review.references
    ])

    latex = f"""\\documentclass{{article}}
\\usepackage[utf8]{{inputenc}}
\\usepackage{{natbib}}

\\title{{{escape_latex(review.title)}}}
\\date{{\\today}}

\\begin{{document}}

\\maketitle

\\section{{Introduction}}
{escape_latex(review.introduction)}

\\section{{Methodology Overview}}
{escape_latex(review.methodology_overview)}

\\section{{Key Findings}}
{escape_latex(review.key_findings)}

\\section{{Research Gaps and Future Directions}}
{escape_latex(review.research_gaps)}

\\section{{Conclusion}}
{escape_latex(review.conclusion)}

\\begin{{thebibliography}}{{99}}
{references_text}
\\end{{thebibliography}}

\\end{{document}}
"""
    return latex


def export_to_word(review: LiteratureReviewResult, output_path: str) -> str:
    """
    Export the literature review to Word (.docx) format.

    Requires python-docx: pip install python-docx
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx is required for Word export. Install with: pip install python-docx")

    doc = Document()

    # Title
    title = doc.add_heading(review.title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Introduction
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(review.introduction)

    # Methodology Overview
    doc.add_heading("Methodology Overview", level=1)
    doc.add_paragraph(review.methodology_overview)

    # Key Findings
    doc.add_heading("Key Findings", level=1)
    doc.add_paragraph(review.key_findings)

    # Research Gaps
    doc.add_heading("Research Gaps and Future Directions", level=1)
    doc.add_paragraph(review.research_gaps)

    # Conclusion
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(review.conclusion)

    # References
    doc.add_heading("References", level=1)
    for ref in review.references:
        doc.add_paragraph(f"{ref['index']}. {ref['formatted']}")

    # Save
    doc.save(output_path)
    return output_path
