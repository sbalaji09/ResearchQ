import re
from typing import Optional
from literature_review_generator import LiteratureReviewResult


# export literature review to markdown format
def export_to_markdown(review_result: LiteratureReviewResult) -> str:
    sections = [
        f"# {review_result.title}",
        "",
        "## Introduction",
        review_result.introduction,
        "",
        "## Methodology Overview",
        review_result.methodology_overview,
        "",
        "## Key Findings",
        review_result.key_findings,
        "",
        "## Research Gaps",
        review_result.research_gaps,
        "",
        "## Conclusion",
        review_result.conclusion,
        "",
        "## References",
    ]

    for i, ref in enumerate(review_result.references, 1):
        authors = ", ".join(ref.get("authors", ["Unknown"]))
        title = ref.get("title", "Untitled")
        year = ref.get("year", "n.d.")
        sections.append(f"{i}. {authors} ({year}). {title}.")

    return "\n".join(sections)


# convert citation numbers to latex cite commands
def _convert_citations_to_latex(text: str, references: list) -> str:
    def replace_citation(match):
        citation_num = int(match.group(1))
        idx = citation_num - 1
        if 0 <= idx < len(references):
            ref = references[idx]
            authors = ref.get("authors", ["unknown"])
            year = ref.get("year", "")
            first_author = authors[0].split()[-1].lower() if authors else "unknown"
            cite_key = f"{first_author}{year}"
            return f"\\cite{{{cite_key}}}"
        return match.group(0)

    return re.sub(r'\[(\d+)\]', replace_citation, text)


# escape special latex characters
def _escape_latex(text: str) -> str:
    replacements = [
        ('\\', '\\textbackslash{}'),
        ('&', '\\&'),
        ('%', '\\%'),
        ('$', '\\$'),
        ('#', '\\#'),
        ('_', '\\_'),
        ('{', '\\{'),
        ('}', '\\}'),
        ('~', '\\textasciitilde{}'),
        ('^', '\\textasciicircum{}'),
    ]
    for old, new in replacements:
        text = text.replace(old, new)
    return text


# export literature review to latex format
def export_to_latex(review_result: LiteratureReviewResult) -> str:
    lines = [
        "\\documentclass{article}",
        "\\usepackage[utf8]{inputenc}",
        "\\usepackage{natbib}",
        "\\usepackage{hyperref}",
        "",
        f"\\title{{{_escape_latex(review_result.title)}}}",
        "\\date{}",
        "",
        "\\begin{document}",
        "\\maketitle",
        "",
    ]

    sections = [
        ("Introduction", review_result.introduction),
        ("Methodology Overview", review_result.methodology_overview),
        ("Key Findings", review_result.key_findings),
        ("Research Gaps", review_result.research_gaps),
        ("Conclusion", review_result.conclusion),
    ]

    for section_title, content in sections:
        converted = _convert_citations_to_latex(content, review_result.references)
        lines.append(f"\\section{{{section_title}}}")
        lines.append(converted)
        lines.append("")

    lines.append("\\bibliographystyle{apalike}")
    lines.append("\\begin{thebibliography}{99}")
    lines.append("")

    for ref in review_result.references:
        authors = ref.get("authors", ["Unknown"])
        title = ref.get("title", "Untitled")
        year = ref.get("year", "n.d.")

        first_author = authors[0].split()[-1].lower() if authors else "unknown"
        cite_key = f"{first_author}{year}"

        author_str = " and ".join(authors)
        lines.append(f"\\bibitem{{{cite_key}}} {_escape_latex(author_str)} ({year}). \\textit{{{_escape_latex(title)}}}.")
        lines.append("")

    lines.append("\\end{thebibliography}")
    lines.append("\\end{document}")

    return "\n".join(lines)


# export literature review to word document
def export_to_word(review_result: LiteratureReviewResult, output_path: Optional[str] = None) -> bytes:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from io import BytesIO

    doc = Document()

    # title
    title_para = doc.add_heading(review_result.title, level=0)
    title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # sections
    sections = [
        ("Introduction", review_result.introduction),
        ("Methodology Overview", review_result.methodology_overview),
        ("Key Findings", review_result.key_findings),
        ("Research Gaps", review_result.research_gaps),
        ("Conclusion", review_result.conclusion),
    ]

    for section_title, content in sections:
        doc.add_heading(section_title, level=1)
        para = doc.add_paragraph(content)
        para.paragraph_format.space_after = Pt(12)

    # references
    doc.add_heading("References", level=1)

    for i, ref in enumerate(review_result.references, 1):
        authors = ", ".join(ref.get("authors", ["Unknown"]))
        title = ref.get("title", "Untitled")
        year = ref.get("year", "n.d.")

        ref_para = doc.add_paragraph()
        ref_para.add_run(f"[{i}] ").bold = True
        ref_para.add_run(f"{authors} ({year}). ")
        ref_para.add_run(title).italic = True
        ref_para.add_run(".")
        ref_para.paragraph_format.left_indent = Inches(0.5)
        ref_para.paragraph_format.first_line_indent = Inches(-0.5)

    if output_path:
        doc.save(output_path)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
